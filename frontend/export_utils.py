from __future__ import annotations

import os
import platform

from docx import Document
from fpdf import FPDF

_FONT_SEARCH_PATHS: dict[str, list[tuple[str, str, str]]] = {
    "Windows": [
        ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf", "C:/Windows/Fonts/ariali.ttf"),
    ],
    "Darwin": [
        ("/Library/Fonts/Arial.ttf", "/Library/Fonts/Arial Bold.ttf", "/Library/Fonts/Arial Italic.ttf"),
        ("/System/Library/Fonts/Supplemental/Arial.ttf",
         "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
         "/System/Library/Fonts/Supplemental/Arial Italic.ttf"),
    ],
    "Linux": [
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf"),
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf"),
    ],
}


def _setup_pdf_fonts(pdf: FPDF) -> str:
    for regular, bold, italic in _FONT_SEARCH_PATHS.get(platform.system(), []):
        if os.path.exists(regular):
            name = "CustomFont"
            pdf.add_font(name, "", regular)
            if os.path.exists(bold):
                pdf.add_font(name, "B", bold)
            if os.path.exists(italic):
                pdf.add_font(name, "I", italic)
            return name
    return "helvetica"


def export_to_txt(title: str, questions: list[dict], filepath: str) -> None:
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write("=" * len(title) + "\n\n")

        for i, q in enumerate(questions, start=1):
            f.write(f"{i}. {q['question']}\n")
            for ans in q.get("answers", []):
                f.write(f"   {ans}\n")
            f.write(f"\nPoprawna odpowiedź: {q.get('correct', '?')}\n")
            if q.get("explanation"):
                f.write(f"Wyjaśnienie: {q['explanation']}\n")
            f.write("\n")


def export_to_pdf(title: str, questions: list[dict], filepath: str) -> None:
    pdf = FPDF()
    pdf.add_page()

    font = _setup_pdf_fonts(pdf)

    pdf.set_font(font, "B", 16)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    for i, q in enumerate(questions, start=1):
        pdf.set_font(font, "B", 12)
        pdf.multi_cell(0, 8, f"{i}. {q['question']}", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font(font, "", 11)
        for ans in q.get("answers", []):
            pdf.multi_cell(0, 6, f"   {ans}", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(2)

        pdf.set_font(font, "B", 11)
        pdf.cell(0, 6, f"Poprawna odpowiedź: {q.get('correct', '?')}", new_x="LMARGIN", new_y="NEXT")

        if q.get("explanation"):
            pdf.set_font(font, "I", 11)
            pdf.multi_cell(0, 6, f"Wyjaśnienie: {q['explanation']}", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(5)

    pdf.output(filepath)


def export_to_docx(title: str, questions: list[dict], filepath: str) -> None:
    doc = Document()
    doc.add_heading(title, 0)

    for i, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {q['question']}").bold = True

        for ans in q.get("answers", []):
            doc.add_paragraph(f"   {ans}")

        p_correct = doc.add_paragraph()
        p_correct.add_run(f"Poprawna odpowiedź: {q.get('correct', '?')}").bold = True

        if q.get("explanation"):
            p_expl = doc.add_paragraph()
            p_expl.add_run(f"Wyjaśnienie: {q['explanation']}").italic = True

    doc.save(filepath)
